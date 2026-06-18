"""
letter_engine.py — Génération de lettres de motivation par LLM
- Ton configurable via config['letter_tone']
- Post-processing : supprime les patterns IA (--, —, clichés)
- Clonage template .docx (formatage préservé)
"""
from src.colors import green as _green, red as _red, yellow as _yellow
from src.colors import cyan as _cyan, bold as _bold, dim as _dim


import re
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
except ImportError:
    Document = None

# ── Tons de lettre ─────────────────────────────────────

TONE_PROMPTS = {
    "professionnel direct": """Le ton doit être professionnel, direct et orienté résultats.
Phrases courtes. Verbes d'action. Chiffres concrets.
Évite les formules de politesse trop longues.""",

    "formel classique": """Le ton doit être formel et classique, comme une lettre traditionnelle.
Formules de politesse élaborées. Vocabulaire soutenu.
Structure académique en 3 parties.""",

    "décontracté": """Le ton doit être décontracté mais professionnel.
Comme un email à un recruteur dans une startup.
Phrases naturelles, tutoiement interdit mais proximité acceptée.""",

    "motivant": """Le ton doit être motivant et montrer de l'enthousiasme concret.
Montre l'envie de rejoindre CETTE entreprise spécifiquement.
Énergie palpable sans être naïf ou exagéré.""",

    "technique": """Le ton doit être technique et précis.
Mets en avant les compétences techniques concrètes.
Termes précis, chiffres, méthodologies.""",
}


# ── Patterns IA à supprimer ────────────────────────────────

IA_CLEANUP_PATTERNS = [
    # Double dash / em dash (signature IA)
    (r'(?<!\w)—(?!\w)', ','),        # em dash isolé → virgule
    (r'(?<!\w)--(?!\w)', ','),       # double dash → virgule
    (r'—\s*—', ','),                 # double em dash
    (r'(\w)—(\w)', r'\1 — \2'),     # em dash entre mots → "mot — mot"
    (r'\s{2,}—\s{2,}', '. '),        # em dash milieu phrase
    # Clichés IA en français
    (r'dans un monde en constante évolution', 'dans notre secteur'),
    (r'au cœur de la transformation numérique', 'dans la transformation de vos processus'),
    (r'à l\'intersection de', 'au croisement de'),
    (r'je suis convaincu(e)?(?: que)?', 'je constate que'),
    (r'j\'ai toujours été passionné(e)? par', 'je m\'intéresse à'),
    (r'(?:tous|chacun) de ces éléments', 'ces éléments'),
    (r'en conclusion', 'pour conclure'),
    # Adjectifs vides IA
    (r'\brobuste\b', 'efficace'),
    (r'\bholistique\b', 'complète'),
    (r'\bsynergique\b', 'coordonnée'),
    (r'\bdisruptif\b', 'innovant'),
    # Phrases trop longues (subordonnées empilées)
    # On ne coupe pas mais on signale
]


def _verify_company_in_body(body, expected_company):
    """Vérifie que l'entreprise cible est mentionnée dans le corps.
    Si absente, tente de l'injecter intelligemment."""
    if not expected_company:
        return body
    if expected_company.lower() in body.lower():
        return body
    # Si l'entreprise n'est pas mentionnée, on l'ajoute dans le closing
    body = body.rstrip()
    body += f"\n\n{expected_company} est exactement le type d'entreprise où je veux investir mon double profil finance et IA."
    return body


def _clean_ia_patterns(text):
    """Supprime les patterns typiques des IA du texte généré."""
    result = text

    # Vérifie et supprime les en-têtes parasites (ex: "SPIE Industrie" en haut)
    result = re.sub(r'^\s*SPIE\s+Industrie\s*\n+', '', result, flags=re.IGNORECASE)
    result = re.sub(r'^\s*SPIE\s+Batignolles\s*\n+', '', result, flags=re.IGNORECASE)

    for pattern, replacement in IA_CLEANUP_PATTERNS:
        result = re.sub(pattern, replacement, result, flags=re.IGNORECASE)

    # Supprime les lignes vides consécutives (max 1)
    result = re.sub(r'\n{3,}', '\n\n', result)

    # Supprime le markdown qui pourrait traîner
    result = result.replace('**', '').replace('*', '')

    return result


# ── Providers LLM ──────────────────────────────────────

def _call_llm(prompt, config):
    """Appelle le LLM configuré pour générer une lettre. Retourne le texte NETTOYÉ."""
    import json
    import ssl
    import urllib.request
    import urllib.error

    llm_config = config.get("llm", {})
    provider = llm_config.get("provider", "none")
    api_config = config.get("api", {})

    if provider == "none":
        return None

    _ssl_ctx = ssl.create_default_context()

    # Prompt système avec ton configuré
    tone = config.get("letter_tone", "professionnel direct")
    tone_instructions = TONE_PROMPTS.get(tone, TONE_PROMPTS["professionnel direct"])

    system_prompt = f"""Tu es un expert en rédaction de lettres de motivation françaises.
Règles ABSOLUES:
1. Rédige en français (niveau C2)
2. Maximum 350 mots
3. {tone_instructions}
4. Structure: 5 paragraphes (accroche, valeur, réalisations, alignement, closing)
5. Utilise des verbes d'action et des chiffres

INTERDITS ABSOLUS (ces patterns sont détectables par les recruteurs):
- JAMAIS "motivé", "dynamique", "passionné" sans preuve concrète
- JAMAIS "Dans un monde en constante évolution" ou variantes
- JAMAIS "À l'intersection de" (cliché IA)
- JAMAIS "Je suis convaincu que" / "J'ai toujours été passionné par"
- JAMAIS "En tant que" en début de phrase (sauf une fois)
- JAMAIS de double tiret (--) ou tiret cadratin (—) comme ponctuation
- JAMAIS d'adjectifs vides: "robuste", "holistique", "synergique", "disruptif"
- JAMAIS de questions rhétoriques ("Comment? En faisant...")
- JAMAIS "Je me permets de vous contacter"
- JAMAIS "C'est avec grand intérêt que" ou variantes

RÈGLES DE STYLE:
- Alterne phrases courtes et longues (rythme naturel)
- Pas plus de 2 phrases avec la même structure consécutive
- Le mot "également" maximum 1 fois dans toute la lettre
- Pas de tirets doubles ou em-dashes -- utilise des virgules ou parenthèses

Réponds UNIQUEMENT par le texte de la lettre. Pas de salutation, pas de markdown."""

    if provider == "ollama":
        model = llm_config.get("model", "llama3.2")
        try:
            data = json.dumps({
                "model": model,
                "prompt": f"{system_prompt}\n\n{prompt}",
                "stream": False,
                "options": {"temperature": 0.7, "num_predict": 800},
            }).encode()
            req = urllib.request.Request(
                "http://localhost:11434/api/generate",
                data=data,
                headers={"Content-Type": "application/json"},
            )
            with urllib.request.urlopen(req, timeout=120) as resp:
                raw = json.loads(resp.read()).get("response", "").strip()
                return _clean_ia_patterns(raw)
        except Exception as e:
            print(f"  " + _yellow(f" Ollama error: {e} "))
            return None

    # Providers OpenAI-compatibles (mistral, openai, deepseek, openrouter, groq, custom)
    PROVIDERS = {
        "openai":       ("openai",      "gpt-5.4-mini",       "https://api.openai.com/v1"),
        "mistral":      ("mistral",     "mistral-small-2506", "https://api.mistral.ai/v1"),
        "deepseek":     ("deepseek",    "deepseek-v4-flash",   "https://api.deepseek.com/v1"),
        "openrouter":   ("openrouter",  "openai/gpt-5.4-mini",  "https://openrouter.ai/api/v1"),
        "groq":         ("groq",        "llama-3.3-70b-versatile", "https://api.groq.com/openai/v1"),
        "custom":       ("custom",      "gpt-5.4-mini",         llm_config.get("base_url", "")),
    }

    if provider in PROVIDERS:
        key_name, default_model, base_url = PROVIDERS[provider]
        key = api_config.get(key_name, {}).get("api_key", "")
        model = llm_config.get("model", default_model)
        if not key or not base_url:
            print(f"  " + _yellow(f" Pas de clé API pour {provider} "))
            return None
        try:
            data = json.dumps({
                "model": model,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt},
                ],
                "temperature": 0.7,
                "max_tokens": 800,
            }).encode()
            req = urllib.request.Request(
                base_url + "/chat/completions",
                data=data,
                headers={
                    "Authorization": f"Bearer {key}",
                    "Content-Type": "application/json",
                },
            )
            with urllib.request.urlopen(req, context=_ssl_ctx, timeout=60) as resp:
                result = json.loads(resp.read())
                raw = result["choices"][0]["message"]["content"].strip()
                return _clean_ia_patterns(raw)
        except Exception as e:
            print(f"  " + _yellow(f" {provider} error: {e} "))
            return None

    return None


# ── Construction du prompt ───────────────────────────────

# Angles de lettre génériques (neutres, pour tout utilisateur)
_LETTER_ANGLES = [
    "le builder qui construit de zéro avec obsession du résultat",
    "le profil qui a vécu la transformation de l'intérieur, pas en consultant",
    "le technicien qui maîtrise la stack et comprend le business",
    "l'expert qui a déployé en production, pas seulement en POC",
    "l'intrapreneur qui sait ce que ça coûte et ce que ça rapporte",
    "le consultant qui prêche par la preuve, pas par les slides",
    "le pont entre deux mondes — crédible des deux côtés",
    "le profil qui build ses propres outils plutôt que d'attendre la DSI",
    "l'approche sécurité d'abord, déploiement en environnement contraint",
    "l'automatisation concrète des processus, avec des chiffres mesurés",
    "l'accélérateur dans une structure en hypercroissance",
    "le conseil qui quantifie le ROI en euros, pas en promesses",
    "la vision stratégique nourrie par l'expérience du terrain",
    "le profil atypique qui ne coche pas toutes les cases mais apporte plus",
]

_LAST_ANGLES = []  # mémoire des 3 derniers angles utilisés

def _pick_angle(offer):
    """Choisit un angle distinct pour cette offre, en évitant les 3 derniers utilisés."""
    import random
    available = [a for a in _LETTER_ANGLES if a not in _LAST_ANGLES]
    if not available:
        available = _LETTER_ANGLES
    angle = random.choice(available)
    _LAST_ANGLES.append(angle)
    if len(_LAST_ANGLES) > 3:
        _LAST_ANGLES.pop(0)
    return angle


def _build_profile_text(profile_data):
    """Construit un texte de profil DYNAMIQUE à partir des données configurées.
    Rien en dur — tout vient du profil utilisateur."""
    if not profile_data:
        return ""
    
    parts = []
    
    # Nom
    name = profile_data.get("name", "")
    if name:
        parts.append(f"CANDIDAT: {name}")
    
    # Titre/poste actuel
    current = profile_data.get("current_position", "")
    if current:
        parts.append(f"POSTE ACTUEL: {current}")
    
    # Expérience
    experience = profile_data.get("experience", "")
    if experience:
        parts.append(f"EXPÉRIENCE: {experience}")
    
    # Formation
    education = profile_data.get("education", "")
    if education:
        parts.append(f"FORMATION: {education}")
    
    # Compétences
    skills = profile_data.get("skills", [])
    if skills:
        parts.append(f"COMPÉTENCES: {', '.join(skills[:8])}")
    
    # Réalisations chiffrées
    metrics = profile_data.get("metrics", [])
    if metrics:
        if isinstance(metrics[0], (list, tuple)):
            metrics_str = ", ".join([f"{v} {u}" for v, u in metrics[:5]])
        else:
            metrics_str = ", ".join(str(m) for m in metrics[:5])
        parts.append(f"RÉALISATIONS: {metrics_str}")
    
    # Disponibilité
    available = profile_data.get("available", "")
    if available:
        parts.append(f"DISPONIBILITÉ: {available}")
    
    # Langues
    languages = profile_data.get("languages", "")
    if languages:
        parts.append(f"LANGUES: {languages}")
    
    # Bio libre (le plus important — l'utilisateur écrit son pitch)
    bio = profile_data.get("bio", "")
    if bio:
        parts.append(f"BIO: {bio}")
    
    return "\n".join(parts)


def _build_letter_prompt(offer, config, profile_data):
    """Construit le prompt pour générer une lettre personnalisée et UNIQUE.
    Tout est dynamique — aucune donnée personnelle en dur."""
    profile = profile_data if profile_data else config.get("profile", {})
    name = profile.get("name", "Le candidat")
    tone = config.get("letter_tone", "professionnel direct")
    angle = _pick_angle(offer)
    
    # Construit le texte de profil dynamiquement
    profile_text = _build_profile_text(profile)
    
    prompt = f"""Génère une lettre de motivation UNIQUE et PERSONNALISÉE.

POSTE: {offer.get('title', '')}
ENTREPRISE: {offer.get('company', '')}
LOCALISATION: {offer.get('location', '')}
DESCRIPTION: {offer.get('description', '')[:500]}

{profile_text}

TON: {tone}

ANGLE IMPOSÉ POUR CETTE LETTRE (ne pas dévier) : {angle}

RÈGLES ABSOLUES:
1. Commence par une phrase d'accroche qui part DU POSTE ou de L'ENTREPRISE, pas du candidat.
2. Varie la structure — pas de pattern répété d'une lettre à l'autre.
3. Mentionne un élément SPÉCIFIQUE à cette offre (produit, chiffre, secteur, enjeu) dans le premier paragraphe.
4. Chaque paragraphe doit commencer différemment.
5. Maximum 350 mots.
6. JAMAIS de clichés: « Je me permets de vous contacter », « À l'intersection de », « Dans un monde en constante évolution ».
7. Pas de tirets doubles (--) ni de tiret cadratin (—).
8. Le mot « également » maximum UNE fois dans toute la lettre.
9. Des verbes d'action concrets: j'ai réduit, j'ai construit, j'ai déployé.
10. Des phrases courtes et longues alternées, rythme naturel.

Réponds UNIQUEMENT par le texte de la lettre. Pas de markdown, pas de titre, pas de signature."""

    return prompt


# ── Template .docx par défaut ───────────────────────────

_DEFAULT_TEMPLATE = Path(__file__).resolve().parent.parent / "templates" / "lettre_template.docx"


def _create_default_template_docx():
    """Charge le template .docx par défaut (cloné du CV Matthias, anonymisé).
    Retourne un objet Document prêt à l'emploi, ou None si introuvable.
    """
    if _DEFAULT_TEMPLATE.exists():
        return Document(str(_DEFAULT_TEMPLATE))
    return None


def _strip_signature(body_text):
    """Retire la signature (Cordialement, Nom) du body_text si présente.
    Car le template .docx a déjà sa propre formule de politesse et signature."""
    import re
    # Patterns de signature à retirer
    for pattern in [
        r'\n*Cordialement,?\s*\n[^\n]+\s*$',
        r'\n*Je vous prie d\'agréer[^\n]*\n[^\n]*\s*$',
        r'\n*Bien cordialement,?\s*\n[^\n]+\s*$',
    ]:
        body_text = re.sub(pattern, '', body_text, flags=re.IGNORECASE)
    return body_text.strip()


# ── Clonage template .docx ───────────────────────────────

def _fill_template_docx(doc, output_path, company, title, body_text, recipient="", profile=None):
    """Remplit un template .docx (fichier ou objet Document) avec les données.
    
    Accepte:
    - str: chemin vers un fichier .docx template
    - Document: objet Document déjà chargé
    
    Placeholders supportés:
    {{ENTREPRISE}} {{POSTE}} {{DATE}} {{CORPS}} {{DESTINATAIRE}} {{PRENOM}} {{EMAIL}}
    """
    if Document is None:
        return False

    try:
        if isinstance(doc, str):
            template_path = doc
            if not Path(template_path).exists():
                return False
            doc = Document(template_path)
        elif not hasattr(doc, 'paragraphs'):
            return False

        date_str = datetime.now().strftime("%d/%m/%Y")
        recipient_text = recipient or f"Responsable du recrutement, {company}"
        
        # Infos profil
        name = (profile or {}).get("name", "").split()[0] if (profile or {}).get("name") else ""
        email = (profile or {}).get("email", "")

        # Remplace tous les placeholders dans tous les paragraphes
        # Ordre important : {{CORPS}} en dernier car il modifie le nombre de paragraphes
        for para in doc.paragraphs:
            full = para.text
            if not full:
                continue
            changed = False
            for ph, val in [
                ("{{ENTREPRISE}}", company),
                ("{{POSTE}}", title),
                ("{{DATE}}", date_str),
                ("{{DESTINATAIRE}}", recipient_text),
                ("{{PRENOM}}", name),
                ("{{EMAIL}}", email),
            ]:
                if ph in full:
                    full = full.replace(ph, val)
                    changed = True
            if changed:
                _set_para_text(para, full)

        # {{CORPS}} traité APRÈS tous les autres (modifie la structure)
        for para in doc.paragraphs:
            if "{{CORPS}}" in para.text:
                # Nettoie le body_text : retire la formule de politesse et signature
                # si elles sont déjà dans le template
                clean_body = _strip_signature(body_text)
                try:
                    _inject_corps(para, para.text, clean_body)
                except Exception:
                    # Fallback: remplacement simple si injection échoue
                    txt = para.text.replace("{{CORPS}}", body_text)
                    _set_para_text(para, txt)
                break

        # Tableaux (pied de page etc)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        txt = para.text
                        for ph, val in [
                            ("{{ENTREPRISE}}", company),
                            ("{{POSTE}}", title),
                            ("{{DATE}}", date_str),
                            ("{{DESTINATAIRE}}", recipient_text),
                            ("{{PRENOM}}", name),
                            ("{{EMAIL}}", email),
                        ]:
                            if ph in txt:
                                txt = txt.replace(ph, val)
                        if "{{CORPS}}" in txt:
                            txt = txt.replace("{{CORPS}}", body_text)
                        if txt != para.text:
                            _set_para_text(para, txt)

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return True

    except Exception as e:
        print(f"  " + _yellow(f" Erreur template .docx : {e} "))
        return False


def _set_para_text(para, text):
    """Remplace le texte d'un paragraphe en préservant le style du premier run."""
    # Sauvegarde le style du premier run
    ref_font = None
    if para.runs:
        ref_font = para.runs[0].font
    # Clear
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
        if ref_font:
            para.runs[0].font.name = ref_font.name
            para.runs[0].font.size = ref_font.size
    else:
        run = para.add_run(text)
        if ref_font:
            run.font.name = ref_font.name
            run.font.size = ref_font.size


def _inject_corps(para, full_text, body_text):
    """Injecte le corps multi-paragraphes à la place de {{CORPS}}.
    Utilise une approche simple : split le placeholder et insère après.
    Retourne le texte restant après {{CORPS}}."""
    before, _, after = full_text.partition("{{CORPS}}")
    before = before.strip()
    after = after.strip()

    # Découpe le corps en paragraphes
    paragraphs = [p.strip() for p in body_text.split("\n\n") if p.strip()]
    if len(paragraphs) <= 1:
        paragraphs = [p.strip() for p in body_text.strip().split("\n") if p.strip()]

    if not paragraphs:
        paragraphs = [body_text]

    # Récupère le parent et l'index pour insérer
    parent = para._element.getparent()
    if parent is None:
        # Paragraphe orphelin — fallback safe: remplace tout le texte
        _set_para_text(para, body_text)
        return ""
    idx = list(parent).index(para._element)

    # 1er paragraphe : before + 1er § du corps
    first_text = (before + " " + paragraphs[0]).strip() if before else paragraphs[0]
    _set_para_text(para, first_text)

    # Insère les paragraphes suivants
    for i, p_text in enumerate(paragraphs[1:], 1):
        new_para = para._element.makeelement(para._element.tag, para._element.attrib)
        # Copie les propriétés de style (pPr)
        pPr = para._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if pPr is not None:
            new_para.insert(0, pPr.__copy__())
        # Ajoute un run avec le texte
        r_elem = new_para.makeelement(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r', {}
        )
        t_elem = r_elem.makeelement(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t', {}
        )
        t_elem.text = p_text
        r_elem.append(t_elem)
        new_para.append(r_elem)
        parent.insert(idx + i, new_para)

    return after  # texte après {{CORPS}}


# ── Fallback: lettre template (sans LLM) ────────────────

def _generate_template_letter(offer, config, profile_data):
    """Génère une lettre basique sans LLM — entièrement dynamique.
    Construit à partir du profil configuré, aucune donnée en dur."""
    profile = profile_data if profile_data else config.get("profile", {})
    name = profile.get("name", "Le candidat")
    first_name = name.split()[0] if name else ""
    company = offer.get("company", "")
    title = offer.get("title", "")
    description = offer.get("description", "")
    
    # Éléments de profil dynamiques
    current_pos = profile.get("current_position", "mon poste actuel")
    experience = profile.get("experience", "mon expérience")
    skills = profile.get("skills", [])
    skills_str = ", ".join(skills[:4]) if skills else "mes compétences"
    education = profile.get("education", "")
    available = profile.get("available", "dès que possible")
    languages = profile.get("languages", "")
    bio = profile.get("bio", "")
    
    # Utilise la bio si dispo, sinon construit un résumé basique
    if bio:
        profile_summary = bio
    elif experience:
        profile_summary = experience
    else:
        profile_summary = f"{current_pos}. {skills_str}."
    
    # Extrait le début de la description pour personnaliser l'accroche
    desc_snippet = description[:120].strip() if description else f"votre approche de {title.lower()}"
    # Tronque proprement au dernier mot
    if len(desc_snippet) >= 120 and ' ' in desc_snippet:
        desc_snippet = desc_snippet[:desc_snippet.rindex(' ')] + "..."
    
    # Construction de la lettre avec les variables du profil
    education_line = f"{education}. " if education else ""
    languages_line = f"{languages}. " if languages else ""
    
    body = f"""Madame, Monsieur,

{company}, c'est {desc_snippet}. Vous cherchez un profil pour {title}. J'arrive avec {experience[:80] if experience else 'une expérience solide'} — et une vraie envie d'impact.

{profile_summary[:200]}. C'est ce bagage que je veux mettre au service de {company}.

{skills_str.capitalize()} — voilà ce qui me permet d'apporter des résultats concrets. J'ai l'habitude des environnements exigeants où l'autonomie et l'initiative comptent autant que la compétence technique.

{education_line}{languages_line}Je suis disponible {available} pour un échange quand vous le souhaitez."""
    
    # Signature avec le prénom seulement si dispo
    if first_name:
        body += f"\n\n{first_name}"
    
    return _clean_ia_patterns(body)

def _save_letter_md(output_path, company, title, body_text, location="Paris"):
    """Sauvegarde la lettre en markdown si pas de template."""
    date_str = datetime.now().strftime("%d/%m/%Y")
    content = f"""{title} — {company}
{location}, le {date_str}

À l'attention du Responsable du Recrutement
{company}

{body_text}
"""
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)


# ── Point d'entrée ─────────────────────────────────────

def generate_all(offers, config, profile_data=None, template_path=None, on_progress=None):
    """Génère les lettres pour toutes les offres fournies.
    Priorité :
    1. Template custom .docx fourni (config['_letter_template_path'])
    2. Template .docx par défaut intégré
    3. Fallback .md si python-docx non disponible
    on_progress(i, total, label) appelé à chaque lettre (optionnel).
    Retourne la liste des lettres générées."""
    if not offers:
        return []

    llm_provider = config.get("llm", {}).get("provider", "none")
    generated = []

    # Résout le template : custom > défaut
    use_template_path = template_path or config.get("_letter_template_path", "")
    use_template_doc = None
    use_template_str = None

    if use_template_path and Path(use_template_path).exists():
        use_template_str = use_template_path
    elif Document is not None:
        # Template par défaut intégré
        try:
            use_template_doc = _create_default_template_docx()
        except Exception:
            pass

    total = len(offers)
    for i, offer in enumerate(offers, 1):
        company = offer.get("company", "")
        title = offer.get("title", "")
        safe_id = offer.get("id", "").replace("/", "-").replace(" ", "_")[:60]
        # Nom de fichier lisible : entreprise_titre-slug
        import unicodedata
        company_slug = re.sub(r'[^a-zA-Z0-9]+', '-', company.lower()).strip('-')
        title_slug = re.sub(r'[^a-zA-Z0-9]+', '-', title.lower()).strip('-')[:50]
        readable_name = f"{company_slug}_{title_slug}" if company_slug and title_slug else safe_id

        # Callback progression
        if on_progress:
            on_progress(i, total, f"{company} — {title[:50]}")

        # Génère le corps (LLM ou template)
        body_text = None
        if llm_provider != "none":
            prompt = _build_letter_prompt(offer, config, profile_data)
            body_text = _call_llm(prompt, config)

        if not body_text:
            body_text = _generate_template_letter(offer, config, profile_data)

        # Vérifie que l'entreprise cible est bien mentionnée
        body_text = _verify_company_in_body(body_text, company)

        lettres_dir = Path(config.get("_letters_dir", "lettres"))
        lettres_dir.mkdir(parents=True, exist_ok=True)

        docx_path = str(lettres_dir / f"{readable_name}.docx")
        md_path = str(lettres_dir / f"{readable_name}.md")

        # Tente .docx d'abord (template custom ou défaut)
        docx_ok = False
        if use_template_str:
            docx_ok = _fill_template_docx(use_template_str, docx_path, company, title, body_text, profile=profile_data)
        elif use_template_doc is not None:
            docx_ok = _fill_template_docx(use_template_doc, docx_path, company, title, body_text, profile=profile_data)

        if docx_ok:
            generated.append({"path": docx_path, "offer_id": offer.get("id"), "format": "docx"})
        else:
            # Fallback .md
            _save_letter_md(md_path, company, title, body_text, location=config.get('preferences', {}).get('location', {}).get('city', 'Paris'))
            generated.append({"path": md_path, "offer_id": offer.get("id"), "format": "md"})

    return generated
